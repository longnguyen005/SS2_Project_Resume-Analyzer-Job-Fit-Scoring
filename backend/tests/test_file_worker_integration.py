from __future__ import annotations

import shutil
import tempfile
from pathlib import Path
from unittest import IsolatedAsyncioTestCase
from uuid import uuid4

import httpx
from fastapi import FastAPI
from sqlalchemy import delete, select, text

from app.core.config import settings
from app.db.session import async_session_factory, engine
from app.file_worker.routes import router as file_worker_router
from app.models import CvUpload, JobDescription, User


class FileWorkerIntegrationTests(IsolatedAsyncioTestCase):
    async def asyncSetUp(self) -> None:
        self.client: httpx.AsyncClient | None = None
        self.user_id = None
        self.cv_upload_id = None
        self._temp_dirs: list[str] = []

        await engine.dispose()

        try:
            async with async_session_factory() as db:
                await db.execute(text("SELECT 1"))

                user = User(
                    email=f"file-worker-{uuid4()}@example.com",
                    hashed_password="hashed-password",
                    full_name="File Worker Test",
                    is_active=True,
                )
                db.add(user)
                await db.flush()

                job_description = JobDescription(
                    user_id=user.id,
                    title="Backend Engineer",
                    description_text="Build backend systems with Python and APIs.",
                )
                db.add(job_description)
                await db.flush()

                cv_upload = CvUpload(
                    user_id=user.id,
                    job_description_id=job_description.id,
                    filename="resume.docx",
                    stored_filename="resume.docx",
                    storage_path="/tmp/resume.docx",
                    storage_key=None,
                    storage_url=None,
                    file_type="docx",
                    file_size_bytes=1024,
                    status="pending",
                    failure_reason=None,
                    failed_stage=None,
                    extracted_text=None,
                )
                db.add(cv_upload)
                await db.commit()

                self.user_id = user.id
                self.cv_upload_id = cv_upload.id
        except Exception as exc:
            self.skipTest(f"Integration database is not available in this environment: {exc}")

        self.client = httpx.AsyncClient(
            transport=httpx.ASGITransport(app=_build_test_app()),
            base_url="http://testserver",
        )

    async def asyncTearDown(self) -> None:
        if self.client is not None:
            await self.client.aclose()

        if self.user_id is not None:
            async with async_session_factory() as db:
                await db.execute(delete(CvUpload).where(CvUpload.user_id == self.user_id))
                await db.execute(delete(JobDescription).where(JobDescription.user_id == self.user_id))
                await db.execute(delete(User).where(User.id == self.user_id))
                await db.commit()

        for temp_dir in self._temp_dirs:
            shutil.rmtree(temp_dir, ignore_errors=True)

        await engine.dispose()

    async def test_file_worker_extract_and_validate_success(self) -> None:
        assert self.client is not None
        headers = {"x-internal-workflow-secret": settings.n8n_internal_shared_secret}
        docx_path = self._create_valid_resume_docx()

        async with async_session_factory() as db:
            cv_upload = await db.scalar(select(CvUpload).where(CvUpload.id == self.cv_upload_id))
            self.assertIsNotNone(cv_upload)
            cv_upload.storage_path = docx_path
            cv_upload.file_type = "docx"
            await db.commit()

        extract_response = await self.client.post(
            f"/api/v1/file-worker/cv/{self.cv_upload_id}/extract",
            headers=headers,
        )
        self.assertEqual(extract_response.status_code, 200)
        extracted_body = extract_response.json()
        self.assertIn("Python, FastAPI, PostgreSQL", extracted_body["resume_text"])

        validate_response = await self.client.post(
            f"/api/v1/file-worker/cv/{self.cv_upload_id}/validate",
            headers=headers,
        )
        self.assertEqual(validate_response.status_code, 200)
        validated_body = validate_response.json()
        self.assertEqual(validated_body["cv_upload_id"], str(self.cv_upload_id))

    async def test_file_worker_extract_failure_marks_extract_stage(self) -> None:
        assert self.client is not None
        headers = {"x-internal-workflow-secret": settings.n8n_internal_shared_secret}

        response = await self.client.post(
            f"/api/v1/file-worker/cv/{self.cv_upload_id}/extract",
            headers=headers,
        )
        self.assertEqual(response.status_code, 422)
        self.assertEqual(response.json()["detail"], "DOCX file was not found for text extraction.")

        async with async_session_factory() as db:
            cv_upload = await db.scalar(select(CvUpload).where(CvUpload.id == self.cv_upload_id))

        self.assertIsNotNone(cv_upload)
        self.assertEqual(cv_upload.status, "failed")
        self.assertEqual(cv_upload.failed_stage, "extract")

    async def test_file_worker_validate_failure_marks_extract_stage(self) -> None:
        assert self.client is not None
        headers = {"x-internal-workflow-secret": settings.n8n_internal_shared_secret}

        async with async_session_factory() as db:
            cv_upload = await db.scalar(select(CvUpload).where(CvUpload.id == self.cv_upload_id))
            self.assertIsNotNone(cv_upload)
            cv_upload.extracted_text = "Weekly admin notes and class discussion summary."
            cv_upload.status = "processing"
            cv_upload.failure_reason = None
            cv_upload.failed_stage = None
            await db.commit()

        response = await self.client.post(
            f"/api/v1/file-worker/cv/{self.cv_upload_id}/validate",
            headers=headers,
        )
        self.assertEqual(response.status_code, 422)
        self.assertIn("does not appear to be a resume or CV", response.json()["detail"])

        async with async_session_factory() as db:
            cv_upload = await db.scalar(select(CvUpload).where(CvUpload.id == self.cv_upload_id))

        self.assertIsNotNone(cv_upload)
        self.assertEqual(cv_upload.status, "failed")
        self.assertEqual(cv_upload.failed_stage, "extract")

    def _create_valid_resume_docx(self) -> str:
        temp_dir = tempfile.mkdtemp(prefix="file-worker-resume-")
        self._temp_dirs.append(temp_dir)
        destination = Path(temp_dir) / "resume.docx"
        _create_resume_docx_content(destination)
        return str(destination)


def _build_test_app() -> FastAPI:
    app = FastAPI()
    app.include_router(file_worker_router, prefix=f"{settings.api_v1_prefix}/file-worker")
    return app


def _create_resume_docx_content(destination: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not available in this environment.") from exc

    document = Document()
    document.add_paragraph("Tran Thi B")
    document.add_paragraph("Email: tranthib@example.com")
    document.add_paragraph("Phone: 0987654321")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("Experience")
    document.add_paragraph("- Built APIs in 2024.")
    document.add_paragraph("- Maintained backend services in 2025.")
    document.add_paragraph("Education")
    document.add_paragraph("Bachelor of Computer Science - 2023")
    document.save(destination)
